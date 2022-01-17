from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()


document = Document()

document.add_picture('profile.jpg', width=Inches(2.0))

# name phone and email details
name = input('What is your name?: ')
speak('Hello ' + name + 'How are you today')
speak('What is your phone number?: ')
phone = input('What is your phone number?: ')

email = input('What is your email?: ')

document.add_paragraph(
    name + ' | ' + str(phone) + ' | ' + email)

# about me
document.add_heading('About me!')
about_me = input('Tell me about yourself: ')
document.add_paragraph(about_me)
# document.add_paragraph(input('Tell me about yourself'))

# work experience
document.add_heading('Work experience: ')
p = document.add_paragraph()

company = input('Enter Company: ')
from_date = input('From_date: ')
to_date = input('To_date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input('Describe your experience at: ' + company + ' ')
p.add_run('Experience: ' + experience_details)

# more experiences
while True:
    has_more_experiences = input('Do you have more experiences?: Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company: ')
        from_date = input('From_date: ')
        to_date = input('To_date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at: ' + company + ' ')
        p.add_run('Experience: ' + experience_details)
    else:
        break

# skills
document.add_heading('Skills: ')
skill = input('Enter skill: ')
r = document.add_paragraph(skill)
r.style = 'List Bullet'

while True:
    has_skills = input('Do you have more skills: Yes or No: ')
    if has_skills.lower() == 'yes':
        skill = input('Enter skill: ')
        r = document.add_paragraph(skill)
        r.style = 'List Bullet'
    else:
        break

# footer

section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = ' CV generated using Amigoscode '

document.save('cv.docx')
